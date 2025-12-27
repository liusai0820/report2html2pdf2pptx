"""
文档解析器 - 解析用户输入的文档数据

@input:  文件路径 (PDF/DOCX/MD/TXT/JSON)
@output: load_document() -> {full_content, pages[], title}
@pos:    数据入口，将各种格式的文档统一转换为可处理的结构

⚠️ 一旦我被更新，务必更新：
   1. 我的头部注释
   2. /src/_FOLDER.md
"""
import json
import os
from typing import List, Dict
from rich.console import Console

console = Console()

class DocumentParser:
    @staticmethod
    def parse_json(file_path: str) -> Dict:
        """解析JSON格式的文档"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return data
    
    @staticmethod
    def parse_markdown(file_path: str) -> Dict:
        """解析Markdown格式的文档（简单实现）"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单按## 分割页面
        pages = []
        sections = content.split('\n## ')
        
        for i, section in enumerate(sections):
            if i == 0 and not section.startswith('## '):
                # 第一个section可能包含标题
                if section.strip():
                    lines = section.strip().split('\n', 1)
                    title = lines[0].replace('# ', '')
                    content = lines[1] if len(lines) > 1 else ''
                    pages.append({'title': title, 'content': content})
            else:
                lines = section.split('\n', 1)
                title = lines[0].strip()
                content = lines[1].strip() if len(lines) > 1 else ''
                pages.append({'title': title, 'content': content})
        
        return {
            'pages': pages,
            'style_guide': '默认样式'
        }
    
    @staticmethod
    def validate_document(data: Dict) -> bool:
        """验证文档数据格式"""
        # 如果有 full_content，说明需要 AI 生成大纲，跳过验证
        if 'full_content' in data:
            return True
        
        if 'pages' not in data:
            console.print("[red]✗ 文档缺少'pages'字段[/red]")
            return False
        
        if not isinstance(data['pages'], list):
            console.print("[red]✗ 'pages'必须是数组[/red]")
            return False
        
        if len(data['pages']) == 0:
            console.print("[red]✗ 'pages'不能为空[/red]")
            return False
        
        for i, page in enumerate(data['pages'], 1):
            if 'title' not in page:
                console.print(f"[yellow]⚠ 第{i}页缺少'title'字段[/yellow]")
            if 'content' not in page:
                console.print(f"[yellow]⚠ 第{i}页缺少'content'字段[/yellow]")
        
        return True
    
    @staticmethod
    def _parse_doc_with_antiword(file_path: str) -> Dict:
        """使用 antiword 工具解析旧版 Word 97-2003 (.doc) 文件"""
        import subprocess
        
        try:
            # 调用 antiword 命令行工具
            result = subprocess.run(
                ['antiword', '-m', 'UTF-8', file_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                error_msg = result.stderr.strip()
                if 'not a Word Document' in error_msg:
                    raise ValueError("文件不是有效的 Word 文档")
                raise ValueError(f"antiword 解析失败: {error_msg}")
            
            content = result.stdout.strip()
            
            if not content:
                raise ValueError("文档内容为空")
            
            # 提取标题（使用文件名）
            doc_title = os.path.splitext(os.path.basename(file_path))[0]
            
            console.print(f"[green]✓[/green] 已使用 antiword 解析旧版 .doc 文件")
            
            return {
                'full_content': content,
                'pages': [],
                'title': doc_title,
                'style_guide': '现代企业风格'
            }
            
        except FileNotFoundError:
            # antiword 未安装
            raise ValueError(
                "无法解析旧版 .doc 文件（antiword 工具未安装）。\n"
                "请将文件另存为 .docx 格式后重试，或联系管理员安装 antiword。"
            )
        except subprocess.TimeoutExpired:
            raise ValueError("文档解析超时，请尝试使用较小的文件。")
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict:
        """解析DOCX格式的文档 - 转为完整Markdown（包含表格和SDT）"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检测文件是否为旧版 .doc 格式（Word 97-2003）
        # 旧版 .doc 文件以 D0 CF 11 E0 开头（OLE Compound Document）
        with open(file_path, 'rb') as f:
            header = f.read(8)
            # OLE 格式的魔数：D0 CF 11 E0 A1 B1 1A E1
            if header[:4] == b'\xd0\xcf\x11\xe0':
                # 尝试使用 antiword 工具解析旧版 .doc 文件
                return DocumentParser._parse_doc_with_antiword(file_path)
        
        try:
            doc = Document(file_path)
        except Exception as e:
            error_msg = str(e)
            if "no relationship of type" in error_msg and "officeDocument" in error_msg:
                raise ValueError(
                    "文档格式不兼容。这可能是因为：\n"
                    "1. 文件是旧版 .doc 格式（Word 97-2003），请另存为 .docx 格式\n"
                    "2. 文件已损坏或不是有效的 Word 文档\n"
                    "3. 文件扩展名错误（例如 PDF 重命名为 .docx）\n\n"
                    "请检查文件并重新上传。"
                )
            raise ValueError(f"无法打开 Word 文档: {error_msg}")
        
        markdown_content = []
        
        # 遍历文档中的所有块级元素（段落、表格、SDT等）
        for element in doc.element.body:
            # 处理段落
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para:
                    text = para.text.strip()
                    if not text:
                        markdown_content.append('')
                        continue
                    
                    # 根据样式转换为Markdown
                    if para.style.name.startswith('Heading 1'):
                        markdown_content.append(f'# {text}')
                    elif para.style.name.startswith('Heading 2'):
                        markdown_content.append(f'## {text}')
                    elif para.style.name.startswith('Heading 3'):
                        markdown_content.append(f'### {text}')
                    elif para.runs and len(para.runs) > 0 and para.runs[0].bold:
                        markdown_content.append(f'**{text}**')
                    else:
                        markdown_content.append(text)
            
            # 处理表格
            elif element.tag.endswith('tbl'):
                markdown_content.append('\n### 表格内容\n')
                for table in doc.tables:
                    if table._element == element:
                        # 提取表格内容
                        for i, row in enumerate(table.rows):
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                            markdown_content.append(row_text)
                            # 在表头后添加分隔符
                            if i == 0:
                                markdown_content.append(' | '.join(['---'] * len(row.cells)))
                        markdown_content.append('')
                        break
            
            # 处理 SDT（Structured Document Tag，如自动生成的目录）
            elif element.tag.endswith('sdt'):
                # 提取 SDT 中的所有文本
                sdt_text_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                sdt_content = ''.join([t.text for t in sdt_text_elements if t.text])
                if sdt_content.strip():
                    markdown_content.append(f'\n### 目录\n{sdt_content}\n')
        
        full_markdown = '\n'.join(markdown_content)
        
        # 提取文档标题：从第一个 Heading 1 或文件名
        doc_title = '文档'
        for line in markdown_content:
            if line.startswith('# '):
                doc_title = line[2:].strip()
                break
        
        # 如果没有找到标题，使用文件名
        if doc_title == '文档':
            doc_title = os.path.splitext(os.path.basename(file_path))[0]
        
        # 返回完整文档，不拆分
        return {
            'full_content': full_markdown,
            'pages': [],  # 空列表，让 AI 自己决定如何拆分
            'title': doc_title,
            'style_guide': '现代企业风格，使用科技蓝和活力橙配色，1280x720分辨率'
        }
    
    @staticmethod
    def parse_docx_old(file_path: str) -> Dict:
        """解析DOCX格式的文档 - 旧版本（按标题拆分）"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        doc = Document(file_path)
        pages = []
        current_page = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题（通常是加粗或使用标题样式）
            if para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold):
                # 如果有当前页面，保存它
                if current_page:
                    pages.append(current_page)
                # 开始新页面
                current_page = {
                    'title': text,
                    'content': ''
                }
            elif current_page:
                # 添加到当前页面内容
                current_page['content'] += text + '\n'
        
        # 添加最后一页
        if current_page:
            pages.append(current_page)
        
        # 如果没有检测到标题，将整个文档作为一页
        if not pages:
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            pages.append({
                'title': '文档内容',
                'content': full_text
            })
        
        return {
            'pages': pages,
            'style_guide': '现代企业风格，使用科技蓝和活力橙配色，1280x720分辨率'
        }
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict:
        """解析 PDF 格式的文档 - 提取文字型 PDF 的文本内容"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("需要安装 pdfplumber: pip install pdfplumber")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        markdown_content = []
        
        with pdfplumber.open(file_path) as pdf:
            console.print(f"[cyan]📄[/cyan] 正在解析 PDF ({len(pdf.pages)} 页)...")
            
            for i, page in enumerate(pdf.pages):
                # 提取文本
                text = page.extract_text()
                if text:
                    # 清理文本
                    lines = text.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            markdown_content.append(line)
                    markdown_content.append('')  # 页面分隔
                
                # 提取表格（如果有）
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        markdown_content.append('\n### 表格内容\n')
                        for j, row in enumerate(table):
                            # 过滤空单元格
                            row_text = ' | '.join([cell.strip() if cell else '' for cell in row])
                            if row_text.replace('|', '').replace(' ', ''):  # 非空行
                                markdown_content.append(row_text)
                                if j == 0:
                                    markdown_content.append(' | '.join(['---'] * len(row)))
                        markdown_content.append('')
        
        full_markdown = '\n'.join(markdown_content)
        
        # 如果没有提取到文本，提示用户
        if not full_markdown.strip():
            console.print("[yellow]⚠ PDF 没有提取到文本，可能是扫描版 PDF。请使用文字型 PDF。[/yellow]")
            raise ValueError("PDF 没有提取到文本内容，可能是扫描版 PDF。请使用文字型 PDF 或将其转换为文字版本。")
        
        # 提取文档标题
        doc_title = os.path.splitext(os.path.basename(file_path))[0]
        
        return {
            'full_content': full_markdown,
            'pages': [],
            'title': doc_title,
            'style_guide': '现代企业风格'
        }
    
    @staticmethod
    def load_document(file_path: str) -> Dict:
        """自动识别并加载文档"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.json':
            data = DocumentParser.parse_json(file_path)
        elif ext in ['.md', '.markdown', '.txt']:
            data = DocumentParser.parse_markdown(file_path)
        elif ext in ['.docx', '.doc']:
            data = DocumentParser.parse_docx(file_path)
        elif ext == '.pdf':
            data = DocumentParser.parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}。支持的格式: .json, .md, .txt, .docx, .pdf")
        
        if not DocumentParser.validate_document(data):
            raise ValueError("文档格式验证失败")
        
        # 显示加载信息
        if 'full_content' in data and data['full_content']:
            content_len = len(data['full_content'])
            console.print(f"[green]✓[/green] 已加载文档: {content_len} 字符")
        elif data.get('pages'):
            console.print(f"[green]✓[/green] 已加载文档: {len(data['pages'])} 页")
        else:
            console.print(f"[yellow]⚠[/yellow] 文档内容为空")
        
        return data
